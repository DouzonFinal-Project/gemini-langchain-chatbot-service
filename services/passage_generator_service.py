# services/exam_generator_service.py
import os
import uuid
import tempfile
import logging
from typing import List, Dict, Optional, AsyncGenerator
from datetime import datetime

from sentence_transformers import SentenceTransformer
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import SystemMessage, HumanMessage

from pymilvus import Collection, FieldSchema, CollectionSchema, DataType, utility

from .gemini_service import stream_generate, generate as gemini_generate

import PyPDF2
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

# 설정 / 모델 초기화 (서비스 레벨에서 관리)
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1000,
    chunk_overlap=200,
    separators=["\n\n", "\n", " "]   # 빈 문자열 제거 권장
)

embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
embedding_dimension = 384

system_prompt = """
당신은 초등~중학생 수준에 맞는 국어 교육용 지문을 작성하는 전문가입니다.
제공된 문서의 관련 내용을 참고하여 사용자의 요청에 맞는 완성도 높은 지문을 생성해주세요.
"""

# --- 파일 텍스트 추출 ---
def extract_text_from_pdf_bytes(file_content: bytes) -> str:
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tf:
            tf.write(file_content)
            temp_path = tf.name
        text = ""
        with open(temp_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for p in reader.pages:
                page_text = p.extract_text() or ""
                text += page_text + "\n"
        os.unlink(temp_path)
        logger.info("PDF 추출 완료, 문자수=%d", len(text))
        return text
    except Exception as e:
        logger.exception("PDF 추출 실패")
        raise

def extract_text_from_docx_bytes(file_content: bytes) -> str:
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tf:
            tf.write(file_content)
            temp_path = tf.name
        doc = DocxDocument(temp_path)
        text = ""
        for p in doc.paragraphs:
            text += p.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        os.unlink(temp_path)
        logger.info("DOCX 추출 완료, 문자수=%d", len(text))
        return text
    except Exception:
        logger.exception("DOCX 추출 실패")
        raise

# --- Milvus 헬퍼 ---
def create_milvus_collection(collection_name: str) -> bool:
    try:
        if utility.has_collection(collection_name):
            c = Collection(collection_name)
            c.drop()
        fields = [
            FieldSchema(name="id", dtype=DataType.VARCHAR, max_length=100, is_primary=True),
            FieldSchema(name="text", dtype=DataType.VARCHAR, max_length=8192),
            FieldSchema(name="embedding", dtype=DataType.FLOAT_VECTOR, dim=embedding_dimension),
            FieldSchema(name="chunk_index", dtype=DataType.INT64),
            FieldSchema(name="metadata", dtype=DataType.VARCHAR, max_length=1000)
        ]
        schema = CollectionSchema(fields=fields, description=f"doc chunks {collection_name}")
        Collection(name=collection_name, schema=schema)
        # 인덱스는 Collection 객체를 재취득 후 생성
        c = Collection(collection_name)
        index_params = {"metric_type":"COSINE", "index_type":"IVF_FLAT", "params":{"nlist":128}}
        c.create_index(field_name="embedding", index_params=index_params)
        logger.info("Milvus 컬렉션 생성: %s", collection_name)
        return True
    except Exception:
        logger.exception("Milvus 컬렉션 생성 실패")
        return False

def insert_chunks_to_milvus(collection_name: str, chunks: List[str]) -> bool:
    try:
        c = Collection(collection_name)
        embeddings = embedding_model.encode(chunks)
        embeddings = [e.tolist() for e in embeddings]
        ids = [str(uuid.uuid4()) for _ in chunks]
        chunk_indices = list(range(len(chunks)))
        metadata_list = ["{}" for _ in chunks]
        entities = [ids, chunks, embeddings, chunk_indices, metadata_list]
        c.insert(entities)
        c.flush()
        logger.info("Milvus 삽입 완료: %d 청크", len(chunks))
        return True
    except Exception:
        logger.exception("Milvus 삽입 실패")
        return False

def search_similar_chunks(collection_name: str, query: str, top_k: int = 5) -> List[str]:
    try:
        if not utility.has_collection(collection_name):
            logger.warning("컬렉션 없음: %s", collection_name)
            return []
        c = Collection(collection_name)
        c.load()
        q_emb = embedding_model.encode([query])[0].tolist()
        search_params = {"metric_type":"COSINE", "params":{"nprobe":10}}
        results = c.search(data=[q_emb], anns_field="embedding", param=search_params, limit=top_k, output_fields=["text","chunk_index","metadata"])
        similar = []
        for r in results[0]:
            similar.append(r.entity.get("text"))
        return similar
    except Exception:
        logger.exception("검색 실패")
        return []

# --- 상위 업무: 업로드 처리(파일 읽고 청크로 나누고 Milvus에 삽입) ---
def process_and_store_document(file_bytes: bytes, filename: str) -> Dict[str, any]:
    ext = filename.split(".")[-1].lower()
    if ext == "pdf":
        text = extract_text_from_pdf_bytes(file_bytes)
    elif ext == "docx":
        text = extract_text_from_docx_bytes(file_bytes)
    else:
        raise ValueError("지원되지 않는 파일 형식")
    if len(text.strip()) < 100:
        raise ValueError("추출된 텍스트가 너무 짧음")
    chunks = text_splitter.split_text(text)
    document_id = f"doc_{uuid.uuid4().hex[:8]}"
    collection_name = f"exam_doc_{document_id}"
    if not create_milvus_collection(collection_name):
        raise RuntimeError("컬렉션 생성 실패")
    if not insert_chunks_to_milvus(collection_name, chunks):
        raise RuntimeError("Milvus 삽입 실패")
    return {
        "document_id": document_id,
        "collection_name": collection_name,
        "chunk_count": len(chunks)
    }

# --- 지문 생성(비스트리밍) ---
async def generate_passage_sync(document_id: str, user_prompt: str) -> Dict[str, any]:
    collection_name = f"exam_doc_{document_id}"
    relevant = search_similar_chunks(collection_name, user_prompt, top_k=5)
    if not relevant:
        raise RuntimeError("참고 문서 내용 없음")
    context = "\n\n".join([f"[참고자료 {i+1}]\n{c}" for i, c in enumerate(relevant)])
    human_prompt = f"""다음은 업로드된 문서에서 추출한 관련 내용입니다:

{context}

---

사용자 요청사항:
{user_prompt}

위 참고 자료를 바탕으로 요청사항에 맞는 국어 지문을 생성해주세요.
다른 설명 없이 생성된 지문만 출력하세요.
"""
    messages = [SystemMessage(content=system_prompt), HumanMessage(content=human_prompt)]
    resp = await gemini_generate(messages)
    return {
        "document_id": document_id,
        "generated_passage": resp["content"],
        "character_count": resp["length"],
        "generation_time": resp["generated_at"],
        "status": "completed"
    }

# --- 지문 생성(스트리밍) ---
async def generate_passage_streaming(document_id: str, user_prompt: str) -> AsyncGenerator[Dict[str, any], None]:
    """
    호출자는 이 async generator를 StreamingResponse에 그대로 넘기면 됩니다.
    각 yield는 dict 형태(예: {'type':'chunk','content':...})로 방출됩니다.
    """
    collection_name = f"exam_doc_{document_id}"
    relevant = search_similar_chunks(collection_name, user_prompt, top_k=5)
    if not relevant:
        yield {"type":"error", "content":"참고 문서 내용 없음", "metadata":{}}
        return
    context = "\n\n".join([f"[참고자료 {i+1}]\n{c}" for i, c in enumerate(relevant)])
    human_prompt = f"""다음은 업로드된 문서에서 추출한 관련 내용입니다:

{context}

---

사용자 요청사항:
{user_prompt}

위 참고 자료를 바탕으로 요청사항에 맞는 국어 지문을 생성해주세요.
다른 설명 없이 생성된 지문만 출력하세요.
"""
    messages = [SystemMessage(content=system_prompt), HumanMessage(content=human_prompt)]
    # 시작
    yield {"type":"start", "content":"지문 생성을 시작합니다...", "metadata":{"document_id":document_id}}
    generated_text = ""
    async for chunk in stream_generate(messages):
        generated_text += chunk
        yield {"type":"chunk", "content": chunk, "metadata":{}}
    yield {"type":"end", "content":"", "metadata":{"document_id":document_id, "character_count":len(generated_text), "generation_time":datetime.now().isoformat()}}
